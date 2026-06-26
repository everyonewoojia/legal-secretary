import os
import datetime


async def export_docx(content: str, title: str = "合同") -> str:
    from docx import Document

    doc = Document()
    doc.add_heading(title, level=1)
    for line in content.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    os.makedirs("exports", exist_ok=True)
    filename = f"exports/{title}_{datetime.datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    doc.save(filename)
    return filename


async def export_pdf(content: str, title: str = "合同") -> str:
    from docx import Document
    from docx2pdf import convert

    docx_path = await export_docx(content, title)
    pdf_path = docx_path.replace(".docx", ".pdf")
    convert(docx_path, pdf_path)
    return pdf_path
